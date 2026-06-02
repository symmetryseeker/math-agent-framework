"""
DocumentEngine — 通用文档生成引擎
==================================
将推导和验证结果渲染为多种格式的文档。

支持格式:
    - .qmd (Quarto — 可渲染为 PDF/HTML/DOCX)
    - .docx (直接生成 Word)
    - .tex (LaTeX)
    - .json (结构化数据)
    - .md (Markdown)

设计原则:
    - 模板驱动: 内容与格式分离
    - 模型无关: 接收任意 Dict 结果，按模板渲染
    - 可扩展: 通过注册新渲染器支持新格式
"""

import json
import os
from datetime import datetime
from typing import Any, Callable, Dict, List, Optional


class DocumentEngine:
    """
    通用文档生成引擎。
    """

    def __init__(self, output_dir: str = "./output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self._renderers: Dict[str, Callable] = {}
        self._register_default_renderers()

    def _register_default_renderers(self):
        self._renderers["json"] = self._render_json
        self._renderers["md"] = self._render_markdown
        self._renderers["qmd"] = self._render_quarto
        self._renderers["docx"] = self._render_docx
        self._renderers["tex"] = self._render_latex

    def register_renderer(self, fmt: str, renderer: Callable):
        """注册自定义渲染器"""
        self._renderers[fmt] = renderer

    def render(
        self,
        data: Dict[str, Any],
        fmt: str = "md",
        filename: Optional[str] = None,
        title: str = "Mathematical Derivation Report",
    ) -> str:
        """
        渲染推导结果为文档。

        Args:
            data: 推导结果字典 (来自 PipelineEngine.collect_results())
            fmt: 输出格式 (json/md/qmd/docx/tex)
            filename: 输出文件名（不含扩展名）
            title: 文档标题

        Returns:
            输出文件路径
        """
        if fmt not in self._renderers:
            raise ValueError(f"Unsupported format: {fmt}. Available: {list(self._renderers)}")

        if filename is None:
            filename = f"derivation_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        path = os.path.join(self.output_dir, f"{filename}.{fmt}")
        content = self._renderers[fmt](data, title)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)

        print(f"  Document saved: {path}")
        return path

    # ── Renderers ──

    def _render_json(self, data: dict, title: str) -> str:
        return json.dumps(data, ensure_ascii=False, indent=2, default=str)

    def _render_markdown(self, data: dict, title: str) -> str:
        lines = [
            f"# {title}",
            f"",
            f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*",
            f"",
        ]

        metadata = data.get("metadata", {})
        if metadata:
            lines.append("## Metadata")
            for k, v in metadata.items():
                lines.append(f"- **{k}**: {v}")
            lines.append("")

        steps = data.get("steps", {})
        if steps:
            lines.append("## Derivation Steps")
            for step_name, step_data in steps.items():
                status = step_data.get("status", "unknown")
                emoji = "✅" if status == "success" else "❌" if status == "failed" else "⏭️"
                lines.append(f"### {emoji} {step_name}")

                output = step_data.get("output", {})
                if isinstance(output, dict):
                    for k, v in output.items():
                        if k in ("verified", "step", "title"):
                            continue
                        if isinstance(v, dict):
                            lines.append(f"**{k}**:")
                            for sk, sv in v.items():
                                if isinstance(sv, str) and len(sv) < 200:
                                    lines.append(f"  - {sk}: {sv}")
                                elif isinstance(sv, (int, float)):
                                    lines.append(f"  - {sk}: {sv}")
                        elif isinstance(v, str) and len(v) < 500:
                            lines.append(f"- **{k}**: {v}")
                        elif isinstance(v, (int, float, bool)):
                            lines.append(f"- **{k}**: {v}")

                error = step_data.get("error")
                if error:
                    lines.append(f"  ⚠️ Error: {error}")
                lines.append("")

        return "\n".join(lines)

    def _render_quarto(self, data: dict, title: str) -> str:
        lines = [
            "---",
            f'title: "{title}"',
            f'date: "{datetime.now().strftime("%Y-%m-%d")}"',
            "format:",
            "  docx:",
            "    toc: true",
            "    number-sections: true",
            "---",
            "",
        ]

        steps = data.get("steps", {})
        for step_name, step_data in steps.items():
            step_title = step_data.get("output", {}).get("title", step_name)
            lines.append(f"## {step_title}")
            lines.append("")

            output = step_data.get("output", {})
            if isinstance(output, dict):
                for k, v in output.items():
                    if k in ("verified", "step", "title"):
                        continue
                    if isinstance(v, str) and len(v) < 1000:
                        lines.append(f"**{k}**: {v}")
                        lines.append("")
                    elif isinstance(v, (int, float)):
                        lines.append(f"- {k} = {v}")
                        lines.append("")
                    elif isinstance(v, dict):
                        lines.append(f"### {k}")
                        for sk, sv in v.items():
                            if isinstance(sv, str) and len(sv) < 500:
                                lines.append(f"- {sk}: {sv}")
                        lines.append("")

        return "\n".join(lines)

    def _render_docx(self, data: dict, title: str) -> str:
        """使用 python-docx 生成 Word 文档"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            doc.styles["Normal"].font.size = Pt(11)

            # Title
            h = doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )

            # Metadata
            metadata = data.get("metadata", {})
            if metadata:
                doc.add_heading("Metadata", level=1)
                for k, v in metadata.items():
                    doc.add_paragraph(f"{k}: {v}", style="List Bullet")

            # Steps
            steps = data.get("steps", {})
            if steps:
                doc.add_heading("Derivation Steps", level=1)
                for step_name, step_data in steps.items():
                    output = step_data.get("output", {})
                    step_title = output.get("title", step_name) if isinstance(output, dict) else step_name
                    doc.add_heading(step_title, level=2)

                    if isinstance(output, dict):
                        for k, v in output.items():
                            if k in ("verified", "step", "title"):
                                continue
                            if isinstance(v, dict):
                                doc.add_paragraph(k, style="List Bullet")
                                for sk, sv in v.items():
                                    if isinstance(sv, str) and len(sv) < 500:
                                        doc.add_paragraph(f"  {sk}: {sv}")
                            elif isinstance(v, str) and len(v) < 1000:
                                doc.add_paragraph(f"{k}: {v}")

            path = os.path.join(self.output_dir, f"{title.replace(' ', '_')}.docx")
            doc.save(path)
            return f"[DOCX saved to {path}]"
        except ImportError:
            return self._render_markdown(data, title)

    def _render_latex(self, data: dict, title: str) -> str:
        lines = [
            "\\documentclass{article}",
            "\\usepackage[UTF8]{ctex}",
            "\\usepackage{amsmath,amssymb}",
            "\\usepackage{hyperref}",
            f"\\title{{{title}}}",
            f"\\date{{{datetime.now().strftime('%Y-%m-%d')}}}",
            "\\begin{document}",
            "\\maketitle",
            "",
        ]

        steps = data.get("steps", {})
        for step_name, step_data in steps.items():
            output = step_data.get("output", {})
            step_title = output.get("title", step_name) if isinstance(output, dict) else step_name
            lines.append(f"\\section{{{step_title}}}")

            if isinstance(output, dict):
                for k, v in output.items():
                    if k in ("verified", "step", "title"):
                        continue
                    if isinstance(v, str):
                        safe_v = v.replace("_", "\\_").replace("&", "\\&")
                        lines.append(f"\\textbf{{{k}}}: {safe_v}\\\\")

        lines.append("\\end{document}")
        return "\n".join(lines)
